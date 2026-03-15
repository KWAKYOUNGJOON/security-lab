#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path
from typing import Any, Dict

from parsers.common import load_json
from report_generator import build_render_context, normalize, validate_json


def build_docx_context(data: Dict[str, Any]) -> Dict[str, Any]:
    return build_render_context(data)


def load_docx_modules():
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("python-docx가 필요합니다. 먼저 `pip install -r requirements.txt`를 실행하세요.") from exc

    return {
        "Document": Document,
        "WD_SECTION": WD_SECTION,
        "WD_STYLE_TYPE": WD_STYLE_TYPE,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Cm": Cm,
        "Pt": Pt,
        "RGBColor": RGBColor,
    }


def ensure_style(document: Any, modules: Dict[str, Any], style_name: str, font_name: str, size: int, bold: bool = False) -> None:
    styles = document.styles
    if style_name in styles:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, modules["WD_STYLE_TYPE"].PARAGRAPH)
    style.font.name = font_name
    style._element.rPr.rFonts.set(modules["qn"]("w:eastAsia"), font_name)
    style.font.size = modules["Pt"](size)
    style.font.bold = bold


def apply_document_styles(document: Any, modules: Dict[str, Any]) -> None:
    ensure_style(document, modules, "CoverTitle", "Malgun Gothic", 24, bold=True)
    ensure_style(document, modules, "CoverMeta", "Malgun Gothic", 11)
    ensure_style(document, modules, "BodyText", "Malgun Gothic", 10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Normal"]:
        style = document.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(modules["qn"]("w:eastAsia"), "Malgun Gothic")
    document.styles["Heading 1"].font.size = modules["Pt"](16)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 1"].font.color.rgb = modules["RGBColor"](31, 45, 61)
    document.styles["Heading 2"].font.size = modules["Pt"](13)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 3"].font.size = modules["Pt"](11)
    document.styles["Heading 3"].font.bold = True
    document.styles["Normal"].font.size = modules["Pt"](10)


def add_page_number(paragraph: Any, modules: Dict[str, Any]) -> None:
    run = paragraph.add_run()
    fld_char_begin = modules["OxmlElement"]("w:fldChar")
    fld_char_begin.set(modules["qn"]("w:fldCharType"), "begin")
    instr_text = modules["OxmlElement"]("w:instrText")
    instr_text.set(modules["qn"]("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = modules["OxmlElement"]("w:fldChar")
    fld_char_end.set(modules["qn"]("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_cover_page(document: Any, modules: Dict[str, Any], report_meta: Dict[str, Any]) -> None:
    paragraph = document.add_paragraph(style="CoverTitle")
    paragraph.alignment = modules["WD_ALIGN_PARAGRAPH"].CENTER
    paragraph.add_run("웹 취약점 진단 보고서")

    document.add_paragraph("")
    document.add_paragraph("")

    info = [
        f"프로젝트명: {report_meta.get('project_name')}",
        f"대상: {report_meta.get('target')}",
        f"작성자: {report_meta.get('author') or '-'}",
        f"고객사/기관: {report_meta.get('customer') or '-'}",
        f"생성 시각: {report_meta.get('generated_at')}",
    ]
    for item in info:
        p = document.add_paragraph(style="CoverMeta")
        p.alignment = modules["WD_ALIGN_PARAGRAPH"].CENTER
        p.add_run(item)

    document.add_page_break()


def add_summary_table(document: Any, modules: Dict[str, Any], summary: Dict[str, Any]) -> None:
    document.add_heading("1. 결과 요약", level=1)
    table = document.add_table(rows=6, cols=2)
    table.alignment = modules["WD_TABLE_ALIGNMENT"].CENTER
    table.style = "Table Grid"
    rows = [
        ("전체 취약점 수", str(summary["total_findings"])),
        ("Critical", str(summary["by_severity"]["critical"])),
        ("High", str(summary["by_severity"]["high"])),
        ("Medium", str(summary["by_severity"]["medium"])),
        ("Low", str(summary["by_severity"]["low"])),
        ("Info", str(summary["by_severity"]["info"])),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value


def add_overview_section(document: Any, report_meta: Dict[str, Any]) -> None:
    document.add_heading("2. 점검 개요", level=1)
    for item in [
        f"프로젝트명: {report_meta.get('project_name')}",
        f"대상: {report_meta.get('target')}",
        f"작성자: {report_meta.get('author') or '-'}",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("3. 점검 범위", level=1)
    for item in report_meta.get("assessment_scope", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("4. 기준", level=1)
    references = report_meta.get("references") or {}
    document.add_paragraph(f"KISA 기준: {references.get('kisa_guide', '-')}", style="List Bullet")
    document.add_paragraph(f"OWASP 기준: {references.get('owasp', '-')}", style="List Bullet")


def add_finding_detail(document: Any, finding: Dict[str, Any]) -> None:
    document.add_heading(f"{finding['id']}. {finding['title']}", level=2)
    overview = [
        f"분류: {finding.get('category')}",
        f"위험도: {(finding.get('severity') or 'info').upper()}",
        f"신뢰도: {(finding.get('confidence') or 'medium').upper()}",
        f"상태: {finding.get('status')}",
        f"자산: {finding.get('asset')}",
        f"URL: {finding.get('url') or '-'}",
        f"파라미터: {finding.get('parameter') or '-'}",
        f"도구: {', '.join(finding.get('tools') or [])}",
    ]
    for item in overview:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("설명", level=3)
    document.add_paragraph(finding.get("description") or "")

    document.add_heading("영향", level=3)
    document.add_paragraph(finding.get("impact") or "")

    document.add_heading("위험도 판단 근거", level=3)
    for item in finding.get("severity_reason") or []:
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("재현 절차", level=3)
    for step in finding.get("reproduction_steps") or []:
        document.add_paragraph(str(step), style="List Number")

    document.add_heading("증거", level=3)
    for evidence in finding.get("evidence") or []:
        suffix = f" ({evidence.get('path')})" if evidence.get("path") else ""
        document.add_paragraph(f"[{evidence.get('type')}] {evidence.get('content')}{suffix}", style="List Bullet")

    document.add_heading("권고 사항", level=3)
    for item in finding.get("recommendation") or []:
        document.add_paragraph(str(item), style="List Bullet")

    refs = finding.get("references") or {}
    document.add_heading("참고 분류", level=3)
    document.add_paragraph(f"OWASP: {', '.join(refs.get('owasp') or ['-'])}", style="List Bullet")
    document.add_paragraph(f"KISA: {', '.join(refs.get('kisa') or ['-'])}", style="List Bullet")


def configure_section(document: Any, modules: Dict[str, Any]) -> None:
    section = document.sections[0]
    section.top_margin = modules["Cm"](2.2)
    section.bottom_margin = modules["Cm"](1.8)
    section.left_margin = modules["Cm"](2.5)
    section.right_margin = modules["Cm"](2.0)
    footer = section.footer.paragraphs[0]
    footer.alignment = modules["WD_ALIGN_PARAGRAPH"].CENTER
    footer.add_run("페이지 ")
    add_page_number(footer, modules)


def render_docx(context: Dict[str, Any], output_path: Path) -> None:
    modules = load_docx_modules()
    document = modules["Document"]()
    apply_document_styles(document, modules)
    configure_section(document, modules)

    add_cover_page(document, modules, context["report_meta"])
    add_summary_table(document, modules, context["summary"])
    add_overview_section(document, context["report_meta"])

    document.add_heading("5. 취약점 상세", level=1)
    for finding in context["findings"]:
        add_finding_detail(document, finding)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="정규화된 findings JSON에서 제출형 DOCX 보고서를 생성합니다.")
    parser.add_argument("--input", required=True, help="정규화된 findings JSON 경로")
    parser.add_argument("--schema", required=True, help="JSON Schema 경로")
    parser.add_argument("--output", required=True, help="출력 DOCX 파일 경로")
    parser.add_argument("--skip-validation", action="store_true", help="스키마 검증 생략")
    args = parser.parse_args()

    data = normalize(load_json(Path(args.input)))
    if not args.skip_validation:
        validate_json(data, Path(args.schema))

    render_docx(build_docx_context(data), Path(args.output))
    print(f"[+] DOCX 보고서 생성 완료: {args.output}")


if __name__ == "__main__":
    main()
