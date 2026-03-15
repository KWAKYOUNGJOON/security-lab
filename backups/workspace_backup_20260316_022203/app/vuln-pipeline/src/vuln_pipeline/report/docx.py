from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


class DocxRenderer:
    def render(self, context: dict, output_path: Path) -> None:
        bundle = context["bundle"]
        profile = context["profile"]
        template = context["template"]
        meta = context["document_meta"]
        document = Document()
        _configure_document(document, bundle.run_id, template["cover_title"])
        _add_cover(document, profile, template, meta)
        _add_info_table(document, template, meta)
        _add_summary(document, context)
        _add_issue_table(document, context)
        _add_remediation_plan(document, context)
        _add_issue_details(document, context)
        _add_appendix(document, context)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)


def _configure_document(document: Document, run_id: str, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10.5)
    document.core_properties.title = title
    document.core_properties.subject = f"Run ID {run_id}"
    _add_page_number(section.footer.paragraphs[0])
    header = section.header.paragraphs[0]
    header.text = f"{title} | {run_id}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_cover(document: Document, profile: dict, template: dict, meta: dict) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(template["cover_title"])
    run.bold = True
    run.font.size = Pt(24)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(template["cover_subtitle"])
    if meta.get("logo_path_optional"):
        logo_note = document.add_paragraph()
        logo_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if Path(meta["logo_path_optional"]).exists():
            logo_note.add_run(f"Logo: {Path(meta['logo_path_optional']).name}")
        else:
            logo_note.add_run("Logo: not provided")
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(f"{meta['client_name']} | {meta['engagement_name']} | {profile['id']}")
    meta_paragraph = document.add_paragraph()
    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_paragraph.add_run(
        f"Version {meta.get('report_version', 'v1.0')} | Delivery {meta.get('delivery_date', '')} | Contact {meta.get('contact_email', '')}"
    )
    if template.get("approval_block"):
        approval = document.add_table(rows=2, cols=2)
        approval.style = "Table Grid"
        approval.rows[0].cells[0].text = "Prepared By"
        approval.rows[0].cells[1].text = meta.get("analyst_name", "")
        approval.rows[1].cells[0].text = "Approved By"
        approval.rows[1].cells[1].text = " / ".join(item for item in [meta.get("approver_name", ""), meta.get("approver_title", "")] if item)
    if meta.get("footer_notice"):
        footer_notice = document.add_paragraph()
        footer_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_notice.add_run(meta["footer_notice"])
    document.add_page_break()


def _add_info_table(document: Document, template: dict, meta: dict) -> None:
    document.add_heading("문서 정보", level=1)
    fields = template.get("info_fields", [])
    table = document.add_table(rows=max(len(fields), 1), cols=2)
    table.style = "Table Grid"
    for idx, field in enumerate(fields):
        table.rows[idx].cells[0].text = field
        table.rows[idx].cells[1].text = meta.get(field, "")


def _add_summary(document: Document, context: dict) -> None:
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(context["executive_summary"])
    if context["comparison_summary"].get("available"):
        diff = context["comparison_summary"]
        document.add_heading("이전 실행 대비 변화", level=2)
        document.add_paragraph(
            f"비교 run_id {diff.get('compared_run_id')} 기준으로 신규 {len(diff.get('new_issues', []))}건, "
            f"해결 {len(diff.get('resolved_issues', []))}건, 변경 {len(diff.get('changed_issues', []))}건입니다."
        )
    document.add_heading("Severity 통계", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    table.rows[0].cells[0].text = "Level"
    table.rows[0].cells[1].text = "Count"
    for level in ["Critical", "High", "Medium", "Low", "Info"]:
        row = table.add_row().cells
        row[0].text = level
        row[1].text = str(context["severity_stats"].get(level, 0))


def _add_issue_table(document: Document, context: dict) -> None:
    document.add_heading("Findings 요약", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Medium Shading 1 Accent 1"
    header = table.rows[0].cells
    for idx, text in enumerate(["Issue ID", "제목", "Severity", "Confidence", "CWE", "자산"]):
        header[idx].text = text
    for item in context["narratives"]:
        issue = item["issue"]
        row = table.add_row().cells
        row[0].text = issue.issue_id
        row[1].text = issue.title
        row[2].text = f"{issue.severity.level} ({issue.severity.score})"
        row[3].text = f"{issue.confidence.level} ({issue.confidence.score})"
        row[4].text = issue.primary_cwe or "-"
        row[5].text = ", ".join(issue.affected_assets)


def _add_remediation_plan(document: Document, context: dict) -> None:
    document.add_heading("Remediation Plan", level=1)
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Issue ID", "Title", "Severity", "Immediate Action", "Structural Fix", "Validation After Fix", "Owner", "Target Due"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row_data in context["remediation_plan"]:
        row = table.add_row().cells
        row[0].text = row_data["issue_id"]
        row[1].text = row_data["title"]
        row[2].text = row_data["severity"]
        row[3].text = row_data["immediate"]
        row[4].text = row_data["structural"]
        row[5].text = row_data["validation"]
        row[6].text = row_data["owner"]
        row[7].text = row_data["due"]


def _add_issue_details(document: Document, context: dict) -> None:
    document.add_heading("상세 Findings", level=1)
    profile = context["profile"]
    template = context["template"]
    headings = template.get("remediation_headings", {})
    for item in context["narratives"]:
        issue = item["issue"]
        narrative = item["narrative"]
        document.add_heading(f"{issue.issue_id}. {issue.title}", level=2)
        _para(document, f"한 줄 요약: {narrative['summary_line']}", True)
        _para(document, f"영향 자산: {', '.join(issue.affected_assets)}", True)
        _para(document, f"Severity / Confidence: {issue.severity.level} ({issue.severity.score}) / {issue.confidence.level} ({issue.confidence.score})", True)
        _para(document, f"CWE / OWASP / KISA: {issue.primary_cwe or 'N/A'} / {', '.join(issue.classification.owasp_top10_2025) or 'N/A'} / {', '.join(issue.classification.kisa_categories) or 'N/A'}", True)
        _para(document, f"개요: {narrative['overview']}", True)
        _para(document, f"왜 문제인지: {narrative['why_it_matters']}", True)
        _para(document, f"증거 요약: {_sanitize_text('; '.join(issue.evidence_summary) or '수집된 대표 증거가 없습니다.', profile)}", True)
        _list_section(document, "재현/확인 포인트", [_sanitize_text(point, profile) for point in narrative["repro_points"]])
        _list_section(document, headings.get("immediate", "즉시 조치"), narrative["quick_fix"])
        _list_section(document, headings.get("structural", "근본 개선"), narrative["structural_fix"])
        _list_section(document, headings.get("validation", "조치 후 확인사항"), narrative["validation_after_fix"])
        _para(document, f"실무 주의사항: {narrative['caution']}", True)
        _para(document, f"Reference Labels: {', '.join(narrative['reference_labels']) or 'N/A'}", True)
        if profile.get("show_analyst_note") and narrative.get("analyst_note"):
            _para(document, f"비고: {narrative['analyst_note']}", True)


def _add_appendix(document: Document, context: dict) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("부록", level=1)
    appendix_sections = context["template"].get("appendix_sections", [])
    if "override" in appendix_sections:
        document.add_heading("Override 적용 내역", level=2)
        for target in context["bundle"].override_summary.get("targets", []) or ["적용된 override 없음"]:
            document.add_paragraph(str(target), style="List Bullet")
    if "false_positive" in appendix_sections:
        document.add_heading("False Positive 목록", level=2)
        if context["bundle"].false_positive_findings:
            for finding in context["bundle"].false_positive_findings:
                document.add_paragraph(f"{finding.finding_id} {finding.title}", style="List Bullet")
        else:
            document.add_paragraph("없음")
    if "suppressed" in appendix_sections:
        document.add_heading("Suppressed / Accepted Risk 목록", level=2)
        if context["appendix"]["suppressed_issues"]:
            for item in context["appendix"]["suppressed_issues"]:
                document.add_paragraph(f"{item['issue_id']} {item['title']} ({item['status']})", style="List Bullet")
        else:
            document.add_paragraph("없음")
    if "artifact_index" in appendix_sections:
        document.add_heading("Artifact 참조 인덱스", level=2)
        for artifact in context["appendix"]["artifact_index"]:
            value = artifact["raw"] if context["profile"].get("show_raw_artifacts") else _mask_path(artifact["redacted"] or artifact["raw"])
            document.add_paragraph(f"{artifact['finding_id']} {artifact['type']} - {value}", style="List Bullet")
    if "decision_trace" in appendix_sections:
        document.add_heading("Decision Trace 요약", level=2)
        for item in context["appendix"]["trace_files"]:
            document.add_paragraph(str(item), style="List Bullet")


def _para(document: Document, text: str, bold_label: bool = False) -> None:
    paragraph = document.add_paragraph()
    if bold_label and ": " in text:
        label, rest = text.split(": ", 1)
        run = paragraph.add_run(label + ": ")
        run.bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)


def _list_section(document: Document, title: str, items: list[str]) -> None:
    _para(document, title + ":", False)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _mask_path(value: str) -> str:
    return value.split("\\")[-1] if value else value


def _sanitize_text(value: str, profile: dict) -> str:
    if profile.get("show_internal_paths"):
        return value
    return re.sub(r"[A-Za-z]:\\[^;\s]+", lambda match: _mask_path(match.group(0)), value)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
